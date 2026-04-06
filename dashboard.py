"""
ForecastIQ — Dashboard (v1.0)
Redesigned UI: planner-first layout, reactive SKU/model switching,
confidence band chart, forward planning table, collapsed model details.

Run:
    streamlit run dashboard.py
"""

from __future__ import annotations

import io
import warnings
from pathlib import Path
from typing import Dict, List, Optional

import numpy as np
import pandas as pd
import plotly.graph_objects as go
import streamlit as st

from core_engine import (
    INDUSTRY_PROFILES,
    MIN_DENSITY_RECORDS,
    classify_abc,
    classify_demand_pattern,
    classify_xyz,
    compute_demand_trend,
    data_quality_check,
    detect_granularity,
    forecast_sku,
    generate_demo_data,
    get_industry_profile,
    horizon_unit_label,
    max_forecast_horizon,
    pareto_analysis,
    _normalize_columns,
)

warnings.filterwarnings("ignore")

# ── Analytics ─────────────────────────────────────────────────────────────────
_ANALYTICS_URL = "https://script.google.com/macros/s/AKfycbx_1HTCDP_GuzqScLcM9au-CsOhioWn5xOoWCKwwX0n0dv5jcSm6IFMJx3SgdLJgY_NQQ/exec"

def _get_session_id() -> str:
    """Generate a random 8-char ID once per browser session."""
    if "session_id" not in st.session_state:
        import uuid
        st.session_state.session_id = str(uuid.uuid4())[:8]
    return st.session_state.session_id

def _log_run(skus_count: int, horizon: int, industry: str, granularity: str, sku_mode: str) -> None:
    """Fire-and-forget usage log. Contains zero user data — counts and settings only."""
    import threading, urllib.request, json as _json
    from datetime import datetime, timezone, timedelta
    PST = timezone(timedelta(hours=-8))
    payload = _json.dumps({
        "timestamp":   datetime.now(PST).strftime("%Y-%m-%d %H:%M:%S PST"),
        "session_id":  _get_session_id(),
        "skus_count":  skus_count,
        "horizon":     horizon,
        "industry":    industry,
        "granularity": granularity,
        "sku_mode":    sku_mode,
    }).encode()
    def _post():
        try:
            req = urllib.request.Request(
                _ANALYTICS_URL,
                data=payload,
                headers={"Content-Type": "application/json"},
                method="POST",
            )
            urllib.request.urlopen(req, timeout=4)
        except Exception:
            pass  # never surface analytics errors to the user
    threading.Thread(target=_post, daemon=True).start()

# ── Palette — Slate Pro ───────────────────────────────────────────────────────
BG      = "#0D1117"
BG2     = "#161B22"
CARD    = "#1C2128"
BORDER  = "#21262D"
BORDER2 = "#30363D"
ACCENT  = "#2F81F7"
BLUE    = "#2F81F7"
PINK    = "#f472b6"
TEXT    = "#F0F6FC"
TEXT2   = "#8B949E"
TEXT3   = "#8B949E"
TEXT4   = "#484F58"
GREEN   = "#3FB950"
AMBER   = "#E3B341"
RED     = "#F85149"
TEAL    = "#79C0FF"

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="ForecastIQ",
    page_icon="📈",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── CSS ───────────────────────────────────────────────────────────────────────
st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600&display=swap');
html, body, [class*="css"] {{
    font-family: 'Inter', -apple-system, sans-serif;
    color: {TEXT};
}}
.stApp {{ background: {BG}; }}
#MainMenu, footer, header {{ visibility: hidden; }}

[data-testid="stSidebar"] {{
    background: {BG2} !important;
    border-right: 0.5px solid {BORDER} !important;
}}
[data-testid="stSidebar"] * {{ color: {TEXT2}; }}

.block-container {{
    padding-top: 1rem !important;
    padding-bottom: 1rem !important;
    max-width: 100% !important;
}}

[data-testid="stMetric"] {{
    background: {BG2};
    border-radius: 8px;
    padding: 10px 14px !important;
}}
[data-testid="stMetric"] label {{
    color: {TEXT3} !important;
    font-size: 0.72rem !important;
    text-transform: uppercase;
    letter-spacing: 0.05em;
}}
[data-testid="stMetric"] [data-testid="stMetricValue"] {{
    color: {TEXT} !important;
    font-size: 1.1rem !important;
    font-weight: 500 !important;
}}
[data-testid="stMetric"] [data-testid="stMetricDelta"] {{
    font-size: 0.8rem !important;
}}

.stButton > button {{
    background: {ACCENT} !important;
    color: #ffffff !important;
    border: none !important;
    border-radius: 8px !important;
    font-weight: 500 !important;
    font-size: 13px !important;
    padding: 10px 16px !important;
    width: 100% !important;
}}
.stButton > button p {{ color: #ffffff !important; }}
.stButton > button:hover {{ opacity: 0.88 !important; }}

.stDownloadButton > button {{
    background: {CARD} !important;
    color: {TEXT2} !important;
    border: 0.5px solid {BORDER2} !important;
    border-radius: 6px !important;
    font-size: 12px !important;
    padding: 5px 12px !important;
}}
.stDownloadButton > button:hover {{
    border-color: {ACCENT} !important;
    color: {BLUE} !important;
}}

div[data-testid="stSelectbox"] > label p,
div[data-testid="stMultiSelect"] > label p,
div[data-testid="stRadio"] > label p,
div[data-testid="stSlider"] > label p,
div[data-testid="stFileUploader"] > label p,
.stCheckbox > label span {{
    color: {TEXT3} !important;
    font-size: 0.78rem !important;
    text-transform: uppercase;
    letter-spacing: 0.05em;
}}

div[data-testid="stMultiSelect"] span[data-baseweb="tag"] {{
    background: {CARD} !important;
    border: 0.5px solid {BORDER2} !important;
    border-radius: 4px !important;
    color: {TEXT2} !important;
    font-size: 11px !important;
}}

div[data-testid="stSlider"] [data-baseweb="slider"] div[role="progressbar"] {{
    background: {ACCENT} !important;
}}
div[data-testid="stSlider"] [role="slider"] {{
    background: {ACCENT} !important;
    border: 2px solid {TEAL} !important;
}}
div[data-testid="stSlider"] [data-testid="stThumbValue"] {{
    color: {BLUE} !important;
    font-weight: 600 !important;
    font-size: 13px !important;
}}

div[data-testid="stFileUploader"] section {{
    border-color: {BORDER} !important;
    background: {BG2} !important;
    border-style: dashed !important;
    border-radius: 8px !important;
}}

.stDataFrame {{ border-radius: 8px; overflow: hidden; }}
hr {{ border-color: {BORDER} !important; }}
.stCaption p {{ color: {TEXT4} !important; font-size: 0.75rem !important; }}
.streamlit-expanderHeader {{ color: {TEXT3} !important; font-size: 0.8rem !important; }}
</style>
""", unsafe_allow_html=True)


# ── Session state init ────────────────────────────────────────────────────────
for key, default in [
    ("df_original", None),
    ("forecast_results", {}),
    ("last_run_key", None),
    ("user_has_selected", False),
]:
    if key not in st.session_state:
        st.session_state[key] = default


# ── Helpers ───────────────────────────────────────────────────────────────────

def _fmt(n) -> str:
    try:
        n = int(n)
    except (TypeError, ValueError):
        return str(n)
    if n >= 1_000_000:
        return f"{n/1_000_000:.1f}M"
    if n >= 1_000:
        return f"{n/1_000:.1f}K"
    return f"{n:,}"


def _read_csv_silent(uploaded_file) -> pd.DataFrame:
    """Load and normalise CSV — suppress all warnings from the parser."""
    df = pd.read_csv(uploaded_file)
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        df = _normalize_columns(df)
    return df


@st.cache_data(show_spinner=False)
def _sample_xlsx_bytes() -> bytes:
    """Build a two-sheet Excel: Read Me guide + Sample Data."""
    import openpyxl
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    rng = np.random.default_rng(7)
    dates = pd.date_range("2023-01-02", periods=104, freq="W")
    cats  = ["Electronics","Apparel","Home & Garden","Sporting Goods","Automotive"]
    locs  = ["Store_A","Store_B","Store_C","Warehouse_1","Warehouse_2"]
    rows  = []
    for i in range(10):
        sku = f"SKU-{i+1:03d}"
        base, trend = 80+i*20, 0.5+i*0.2
        for t, d in enumerate(dates):
            s = (10+i*3)*np.sin(2*np.pi*t/52)
            qty = max(base+trend*t+s+rng.normal(0,6+i),0)
            rows.append([d.strftime("%Y-%m-%d"),sku,round(qty,0),cats[i%5],locs[i%5],round(9.99+i*2.5+rng.uniform(-1,1),2)])
    sample_df = pd.DataFrame(rows, columns=["Date","SKU","Quantity","Category","Location","Price"])

    wb = Workbook()
    thin   = Side(style="thin", color="CBD5E1")
    bdr    = Border(left=thin,right=thin,top=thin,bottom=thin)
    DARK   = PatternFill("solid", start_color="1F3864")
    MID    = PatternFill("solid", start_color="2563EB")
    ALT    = PatternFill("solid", start_color="F8FAFC")
    REQ    = PatternFill("solid", start_color="FEF9C3")
    OPT    = PatternFill("solid", start_color="DCFCE7")
    CENT   = Alignment(horizontal="center",vertical="center",wrap_text=True)
    LEFT   = Alignment(horizontal="left",  vertical="center",wrap_text=True)

    def _h(ws,r,c,v,fill=None,bold=True,color="FFFFFF",align=CENT):
        cell=ws.cell(row=r,column=c,value=v)
        cell.font=Font(name="Arial",bold=bold,size=10,color=color)
        cell.fill=fill or DARK; cell.alignment=align; cell.border=bdr
        return cell
    def _c(ws,r,c,v,fill=None,bold=False,align=LEFT):
        cell=ws.cell(row=r,column=c,value=v)
        cell.font=Font(name="Arial",bold=bold,size=10)
        if fill: cell.fill=fill
        cell.alignment=align; cell.border=bdr
        return cell

    # ── Sheet 1: Read Me ──────────────────────────────────────
    ws1=wb.active; ws1.title="Read Me"
    ws1.merge_cells("A1:G1")
    t=ws1["A1"]; t.value="ForecastIQ — Data Format & Output Guide"
    t.font=Font(name="Arial",bold=True,size=13,color="FFFFFF")
    t.fill=DARK; t.alignment=CENT; ws1.row_dimensions[1].height=30

    # Column reference
    ws1.merge_cells("A3:G3")
    s=ws1["A3"]; s.value="REQUIRED & OPTIONAL COLUMNS"
    s.font=Font(name="Arial",bold=True,size=10,color="FFFFFF"); s.fill=MID; s.alignment=CENT
    for ci,h in enumerate(["Column","Required?","Data Type","Format / Example","Accepted Aliases","Notes","Status"],1):
        _h(ws1,4,ci,h)
    col_data=[
        ["Date","Yes","Date (text)","2024-01-07","date, week, order_date, trans_date, period","ISO format preferred. Day-first fallback supported.","Required"],
        ["SKU","Yes","Text","SKU-001 or 1025","sku, item, product, product_code, article","Any alphanumeric ID. Spaces trimmed.","Required"],
        ["Quantity","Yes","Number","120 or 34.5","quantity, qty, sales, demand, units, orders","Must be numeric. $ symbols stripped. Negatives dropped.","Required"],
        ["Category","No","Text","Electronics","category, cat, segment, family, group","Enables category filter in sidebar.","Optional"],
        ["Location","No","Text","Store_A","location, store, warehouse, site, region","Enables location filter.","Optional"],
        ["Price","No","Number","9.99","price, unit_price, checkout_price, cost","Enables revenue-based Pareto ranking.","Optional"],
    ]
    sf={"Required":REQ,"Optional":OPT}
    for ri,row in enumerate(col_data,5):
        bg=ALT if ri%2==0 else None
        for ci,val in enumerate(row,1):
            f=sf.get(val,bg); c=_c(ws1,ri,ci,val,fill=f)
            if ci==2: c.alignment=CENT
            if ci==7:
                c.font=Font(name="Arial",size=10,bold=True,
                            color="92400E" if val=="Required" else "14532D")

    # Limitations
    r=13; ws1.merge_cells(f"A{r}:G{r}")
    s2=ws1[f"A{r}"]; s2.value="DATA REQUIREMENTS & LIMITATIONS"
    s2.font=Font(name="Arial",bold=True,size=10,color="FFFFFF"); s2.fill=MID; s2.alignment=CENT
    lims=[
        ["Min rows per SKU","30 rows minimum. SKUs below this are skipped automatically."],
        ["Recommended history","52+ weeks (1 year) for seasonality. 2+ years is ideal."],
        ["File format","CSV only (.csv). Comma-delimited. UTF-8 encoding."],
        ["Max file size","200 MB per upload."],
        ["Duplicate rows","Duplicate Date+SKU rows are summed before forecasting."],
        ["Zero demand","Allowed. Handled by Croston & TSB models."],
        ["Negative quantity","Dropped during parsing."],
        ["Column order","Does not matter — detection is by name, not position."],
    ]
    ws1.merge_cells(f"B{r+1}:G{r+1}")
    _h(ws1,r+1,1,"Parameter"); _h(ws1,r+1,2,"Detail")
    for ri2,(p,d) in enumerate(lims,r+2):
        bg=ALT if ri2%2==0 else None
        _c(ws1,ri2,1,p,fill=bg,bold=True)
        ws1.merge_cells(f"B{ri2}:G{ri2}"); _c(ws1,ri2,2,d,fill=bg)

    # Output guide
    r2=r+len(lims)+3; ws1.merge_cells(f"A{r2}:G{r2}")
    s3=ws1[f"A{r2}"]; s3.value="WHAT TO EXPECT FROM THE OUTPUT"
    s3.font=Font(name="Arial",bold=True,size=10,color="FFFFFF"); s3.fill=MID; s3.alignment=CENT
    outs=[
        ["Forecast CSV columns","Date, SKU, Forecasted_Qty, Lower_Bound, Upper_Bound, Best_Model, MAPE_pct, Confidence, Confidence_Score, Demand_Pattern"],
        ["Models compared","5 models per SKU: Moving Average, Holt-Winters, Trend+Seasonal, Croston, TSB. Best = lowest WAPE."],
        ["Confidence band","±15% around the active forecast line. Shown as shaded area on the chart."],
        ["Demand patterns","Stable (CV≤0.4), Mixed (0.4–0.8), Volatile (CV>0.8), Intermittent (>50% zero periods)."],
        ["Confidence label","High (score≥75), Medium (50–74), Low (<50). Based on MAPE + cross-validation stability."],
        ["Forward table","3-period quantities per SKU + period total + YoY % vs same period prior year."],
        ["Skipped SKUs","SKUs with <30 data points are skipped and listed with a reason."],
    ]
    ws1.merge_cells(f"B{r2+1}:G{r2+1}")
    _h(ws1,r2+1,1,"Output"); _h(ws1,r2+1,2,"Description")
    for ri3,(o,d) in enumerate(outs,r2+2):
        bg=ALT if ri3%2==0 else None
        _c(ws1,ri3,1,o,fill=bg,bold=True)
        ws1.merge_cells(f"B{ri3}:G{ri3}"); _c(ws1,ri3,2,d,fill=bg)

    for col,w in zip("ABCDEFG",[22,12,18,22,30,36,11]):
        ws1.column_dimensions[col].width=w
    ws1.freeze_panes="A5"

    # ── Sheet 2: Sample Data ──────────────────────────────────
    ws2=wb.create_sheet("Sample Data")
    for ci,h in enumerate(sample_df.columns,1):
        _h(ws2,1,ci,h)
    for ri,row in enumerate(sample_df.itertuples(index=False),2):
        bg=ALT if ri%2==0 else None
        for ci,val in enumerate(row,1):
            c=_c(ws2,ri,ci,val,fill=bg)
            if ci==1: c.alignment=CENT
            if ci==3: c.number_format="#,##0"
            if ci==6: c.number_format="#,##0.00"
    for col,w in zip("ABCDEF",[14,12,12,16,14,10]):
        ws2.column_dimensions[col].width=w
    ws2.freeze_panes="A2"
    ws2.auto_filter.ref=f"A1:F{len(sample_df)+1}"

    buf=io.BytesIO(); wb.save(buf); return buf.getvalue()


@st.cache_data(show_spinner=False)
def _starter_kit_zip() -> bytes:
    """Zip: forecastiq_guide.docx + forecastiq_sample.csv"""
    import zipfile, io as _io
    from docx import Document as _Doc
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _add_shading(cell, fill_hex):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), fill_hex)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    doc = _Doc()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = t.add_run('ForecastIQ — Quick Start Guide')
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x1F,0x38,0x64)

    doc.add_paragraph('Upload your demand data and get an S&OP-ready forecast in under a minute.\nLive tool: preetpatel-forecastiq.streamlit.app — your data never leaves your browser.')

    # Helper: section heading
    def sec(txt):
        p = doc.add_paragraph()
        r2 = p.add_run(txt); r2.bold = True
        r2.font.size = Pt(13); r2.font.color.rgb = RGBColor(0x25,0x63,0xEB)
        doc.add_paragraph()

    # 1 - Required columns
    sec('1  Required columns')
    tbl1 = doc.add_table(rows=1, cols=4)
    tbl1.style = 'Table Grid'
    for i, h in enumerate(['Column','Data type','Example','Accepted names']):
        c = tbl1.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        _add_shading(c, '1F3864')
    for row_data in [
        ['Date',     'Date or text', '2024-01-07',  'date, week, order_date, period'],
        ['SKU',      'Text',         'SKU-001',      'sku, item, product, product_code'],
        ['Quantity', 'Number',       '120 or 34.5',  'quantity, qty, sales, demand, units'],
    ]:
        row = tbl1.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            if i == 0:
                row.cells[i].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    sec('2  Optional columns')
    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = 'Table Grid'
    for i, h in enumerate(['Column','Example','What it unlocks']):
        c = tbl2.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        _add_shading(c, '1F3864')
    for row_data in [
        ['Category', 'Electronics', 'Category filter in sidebar'],
        ['Location', 'Store_A',     'Location filter in sidebar'],
        ['Price',    '9.99',        'Revenue-based Pareto ranking'],
    ]:
        row = tbl2.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            if i == 0:
                row.cells[i].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    r3 = p.add_run('Column order does not matter — detection is by name, not position.')
    r3.italic = True; r3.font.color.rgb = RGBColor(0x37,0x41,0x51)

    doc.add_paragraph()
    sec('3  Data requirements')
    reqs = [
        ('File format',         'CSV only (.csv). Comma-delimited. UTF-8 encoding.'),
        ('Minimum rows',        '30 rows per SKU. SKUs below this are skipped automatically.'),
        ('Recommended history', '52+ weeks (1 year) minimum. 2+ years is ideal.'),
        ('Duplicates',          'Duplicate Date + SKU rows are summed automatically.'),
        ('Zero demand',         'Allowed — handled by Croston and TSB models.'),
        ('Negative quantity',   'Rows with negative quantity are dropped.'),
        ('Max file size',       '200 MB per upload.'),
    ]
    for label, detail in reqs:
        p2 = doc.add_paragraph(style='List Bullet')
        r4 = p2.add_run(f'{label}: '); r4.bold = True
        p2.add_run(detail)

    doc.add_paragraph()
    sec('4  What you get')
    outputs = [
        'Forecast chart — 5 model lines, best model highlighted, confidence band shaded.',
        'Forward planning table — 3-period quantities, quarter total, YoY %, confidence rating.',
        'Export CSV — Date, SKU, Forecasted Qty, Lower/Upper Bound, Best Model, MAPE %, Confidence, Demand Pattern.',
        'Model details — MAE, WAPE, RMSE for all 5 models, collapsed by default.',
    ]
    for o in outputs:
        doc.add_paragraph(o, style='List Bullet')

    doc.add_paragraph()
    sec('5  Example data row')
    tbl3 = doc.add_table(rows=2, cols=6)
    tbl3.style = 'Table Grid'
    for i, h in enumerate(['Date','SKU','Quantity','Category','Location','Price']):
        c = tbl3.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        _add_shading(c, '1F3864')
    for i, v in enumerate(['2024-01-07','SKU-001','120','Electronics','Store_A','9.99']):
        tbl3.rows[1].cells[i].text = v

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    r5 = p3.add_run('The included forecastiq_sample.csv is formatted correctly and ready to upload directly.')
    r5.italic = True; r5.font.color.rgb = RGBColor(0x92,0x40,0x0E)

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    r6 = p4.add_run('Built by Preet Patel  |  M.S. Engineering Management, USC Viterbi \'26  |  preetpatel-forecastiq.streamlit.app')
    r6.font.size = Pt(9); r6.font.color.rgb = RGBColor(0x9C,0xA3,0xAF); r6.italic = True

    docx_buf = _io.BytesIO(); doc.save(docx_buf); docx_bytes = docx_buf.getvalue()

    # Sample CSV
    rng = np.random.default_rng(7)
    dates = pd.date_range("2023-01-02", periods=104, freq="W")
    cats = ["Electronics","Apparel","Home & Garden","Sporting Goods","Automotive"]
    locs = ["Store_A","Store_B","Store_C","Warehouse_1","Warehouse_2"]
    rows = []
    for i in range(10):
        sku = f"SKU-{i+1:03d}"
        base, trend = 80+i*20, 0.5+i*0.2
        for t, d in enumerate(dates):
            s = (10+i*3)*np.sin(2*np.pi*t/52)
            qty = max(base+trend*t+s+rng.normal(0,6+i), 0)
            rows.append([d.strftime("%Y-%m-%d"), sku, round(qty,0), cats[i%5], locs[i%5]])
    csv_bytes = pd.DataFrame(rows, columns=["Date","SKU","Quantity","Category","Location"]
                             ).to_csv(index=False).encode()

    buf = _io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("forecastiq_guide.docx", docx_bytes)
        zf.writestr("forecastiq_sample.csv",  csv_bytes)
    return buf.getvalue()


def _run_key(skus, horizon, industry, location) -> str:
    return f"{sorted(skus)}|{horizon}|{industry}|{location}"


def _build_chart(result: dict, active_model: Optional[str] = None) -> go.Figure:
    hist = result["historical_df"]
    fdf  = result["forecast_df"]
    best = result["best_method"]
    show = active_model or best

    fc_len = len(fdf)
    trail  = min(len(hist), max(fc_len * 3, 52))
    h_tail = hist.tail(trail)

    model_cols = {
        "Moving Average":  "Moving_Average",
        "Holt-Winters":    "Holt_Winters",
        "Trend+Seasonal":  "Trend_Seasonal",
        "Croston":         "Croston",
        "TSB":             "TSB",
    }

    # Confidence band: ±15% around the active model forecast
    active_col = model_cols.get(show, "Best_Model")
    if active_col not in fdf.columns:
        active_col = "Best_Model"
    fc_vals = fdf[active_col].values.astype(float)
    lo = np.maximum(fc_vals * 0.85, 0)
    hi = fc_vals * 1.15

    fig = go.Figure()

    # Confidence band (filled area)
    fig.add_trace(go.Scatter(
        x=pd.concat([fdf["Date"], fdf["Date"][::-1]]),
        y=np.concatenate([hi, lo[::-1]]),
        fill="toself",
        fillcolor="rgba(37,99,235,0.12)",
        line=dict(width=0),
        name="Confidence band",
        hoverinfo="skip",
    ))

    # Historical
    fig.add_trace(go.Scatter(
        x=h_tail["Date"], y=h_tail["Quantity"],
        name="Historical",
        mode="lines",
        line=dict(color=TEXT2, width=1.5),
    ))

    # Alt model styles — each visually distinct
    ALT_STYLES = {
        "Moving Average":  dict(color="#9ca3af", dash="solid",    width=1.2),
        "Holt-Winters":    dict(color="#2dd4bf", dash="dash",     width=1.2),
        "Trend+Seasonal":  dict(color="#a78bfa", dash="dot",      width=1.2),
        "Croston":         dict(color=PINK,      dash="dashdot",  width=1.2),
        "TSB":             dict(color="#fb923c", dash="longdash", width=1.2),
    }

    # All non-active models — distinct colors, reduced opacity
    for label, col in model_cols.items():
        if col not in fdf.columns:
            continue
        if label == show:
            continue
        style = ALT_STYLES.get(label, dict(color=BLUE, dash="dot", width=1))
        fig.add_trace(go.Scatter(
            x=fdf["Date"], y=fdf[col],
            name=label,
            mode="lines",
            line=dict(color=style["color"], width=style["width"], dash=style["dash"]),
            opacity=0.45,
        ))

    # Active model — prominent
    fig.add_trace(go.Scatter(
        x=fdf["Date"], y=fdf[active_col],
        name=f"{show} {'★' if show == best else ''}".strip(),
        mode="lines+markers",
        line=dict(color=PINK, width=2.5, dash="dash"),
        marker=dict(size=5, color=PINK),
    ))

    # Forecast starts line
    last_hist = hist["Date"].iloc[-1]
    fig.add_vline(x=last_hist, line_dash="dash", line_color=TEXT4, line_width=1)
    fig.add_annotation(
        x=last_hist, y=1.04, yref="paper",
        text="Forecast starts", showarrow=False,
        font=dict(size=10, color=TEXT3),
    )

    fig.update_layout(
        template="plotly_dark",
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        height=300,
        margin=dict(l=40, r=10, t=20, b=35),
        legend=dict(
            orientation="h", yanchor="bottom", y=1.06,
            xanchor="right", x=1,
            font=dict(size=10, color=TEXT3),
        ),
        font=dict(family="Inter, sans-serif", color=TEXT3, size=11),
        xaxis=dict(gridcolor="rgba(255,255,255,0.04)", tickfont=dict(size=10, color=TEXT4)),
        yaxis=dict(
            gridcolor="rgba(255,255,255,0.04)",
            tickfont=dict(size=10, color=TEXT4),
            tickformat=",",
            title=dict(text="Quantity", font=dict(size=11, color=TEXT3)),
        ),
    )
    return fig


def _build_export(results: dict, pareto_table: pd.DataFrame = None) -> pd.DataFrame:
    """
    Export CSV: one row per forecast period per SKU.
    Columns: Date, SKU, ABC, XYZ, Forecasted_Qty, Lower_Bound, Upper_Bound,
             Best_Model, WAPE_pct, Confidence
    """
    # Build ABC lookup
    abc_lkp = {}
    if pareto_table is not None and "ABC" in pareto_table.columns:
        abc_lkp = dict(zip(pareto_table["SKU"], pareto_table["ABC"]))

    rows = []
    col_map = {
        "Moving Average": "Moving_Average",
        "Holt-Winters":   "Holt_Winters",
        "Trend+Seasonal": "Trend_Seasonal",
        "Croston":        "Croston",
        "TSB":            "TSB",
    }
    for sku, r in results.items():
        if r.get("skipped"):
            continue
        fdf         = r["forecast_df"].copy()
        best_method = r["best_method"]
        conf        = r.get("confidence", {})
        best_col    = col_map.get(best_method, "Best_Model")
        if best_col not in fdf.columns:
            best_col = "Best_Model"
        fc_vals = fdf[best_col].fillna(0).values.astype(float)
        lo_vals = np.maximum(fc_vals * 0.85, 0).astype(int)
        hi_vals = (fc_vals * 1.15).astype(int)

        # XYZ from post-aggregation CV (correct — engine already computed it)
        cv  = r["demand_profile"].get("cv", 0)
        xyz = classify_xyz(cv if np.isfinite(cv) else 999)
        abc = abc_lkp.get(sku, "")

        export = pd.DataFrame({
            "Date":           pd.to_datetime(fdf["Date"]).dt.strftime("%Y-%m-%d"),
            "SKU":            sku,
            "ABC":            abc,
            "XYZ":            xyz,
            "Forecasted_Qty": fc_vals.astype(int),
            "Lower_Bound":    lo_vals,
            "Upper_Bound":    hi_vals,
            "Best_Model":     best_method,
            "WAPE_pct":       r["accuracy_results"][r["best_index"]]["MAPE_%"],
            "Confidence":     conf.get("label", ""),
        })
        rows.append(export)
    if not rows:
        return pd.DataFrame()
    out = pd.concat(rows, ignore_index=True)
    for col in ["Forecasted_Qty", "Lower_Bound", "Upper_Bound"]:
        out[col] = out[col].apply(lambda x: f"{int(x):,}")
    return out
    if not rows:
        return pd.DataFrame()
    out = pd.concat(rows, ignore_index=True)
    # Comma-format quantity columns as strings for readability
    for col in ["Forecasted_Qty", "Lower_Bound", "Upper_Bound"]:
        out[col] = out[col].apply(lambda x: f"{int(x):,}")
    return out


def _forward_table(results: dict, granularity: str, pareto_table: pd.DataFrame = None) -> pd.DataFrame:
    """
    3-period forward planning table with ABC, XYZ, Trend columns.
    Quantities comma-formatted. No None values.
    """
    col_map = {
        "Moving Average": "Moving_Average",
        "Holt-Winters":   "Holt_Winters",
        "Trend+Seasonal": "Trend_Seasonal",
        "Croston":        "Croston",
        "TSB":            "TSB",
    }

    # ABC lookup from pareto table
    abc_lkp = {}
    if pareto_table is not None and "ABC" in pareto_table.columns:
        abc_lkp = dict(zip(pareto_table["SKU"], pareto_table["ABC"]))

    # Collect period labels
    period_labels = []
    for r in results.values():
        if not r.get("skipped"):
            fdf = r["forecast_df"]
            dates = pd.to_datetime(fdf["Date"].values[:3])
            if granularity == "Monthly":
                period_labels = [d.strftime("%b %Y") for d in dates]
            else:
                period_labels = [d.strftime("%d %b") for d in dates]
            break
    while len(period_labels) < 3:
        period_labels.append(f"P{len(period_labels)+1}")
    period_labels = period_labels[:3]

    rows = []
    for sku, r in results.items():
        if r.get("skipped"):
            continue
        fdf      = r["forecast_df"]
        conf     = r.get("confidence", {})
        best_col = col_map.get(r["best_method"], "Best_Model")
        if best_col not in fdf.columns:
            best_col = "Best_Model"

        fc_series = fdf[best_col].fillna(0).values[:3].astype(float)
        n_periods = len(fc_series)

        # Demand Trend — slope vs recent history, works on any history length
        hist_vals = r["historical_df"]["Quantity"].values.astype(float)
        fc_all    = fdf[best_col].fillna(0).values.astype(float)
        trend_pct = compute_demand_trend(hist_vals, fc_all)
        trend_str = f"+{trend_pct:.1f}%" if trend_pct >= 0 else f"{trend_pct:.1f}%"

        # XYZ from CV
        cv  = r["demand_profile"].get("cv", 0)
        xyz = classify_xyz(cv if np.isfinite(cv) else 999)
        abc = abc_lkp.get(sku, "")

        row = {
            "SKU":            sku,
            "ABC":            abc,
            "XYZ":            xyz,
            "Pattern":        r["demand_profile"]["pattern"],
            period_labels[0]: f"{int(fc_series[0]):,}" if n_periods > 0 else "—",
            period_labels[1]: f"{int(fc_series[1]):,}" if n_periods > 1 else "—",
            period_labels[2]: f"{int(fc_series[2]):,}" if n_periods > 2 else "—",
            "3-period total": f"{int(fc_series.sum()):,}",
            "Trend":          trend_str,
            "Confidence":     conf.get("label", "?"),
        }
        rows.append(row)
    return pd.DataFrame(rows) if rows else pd.DataFrame()


# ── Sidebar ───────────────────────────────────────────────────────────────────
with st.sidebar:
    # Logo
    st.markdown(f"""
    <div style="display:flex;align-items:center;gap:9px;padding:4px 0 12px;">
        <div style="width:26px;height:26px;border-radius:6px;background:{ACCENT};
                    display:flex;align-items:center;justify-content:center;
                    font-size:13px;font-weight:600;color:#fff;">F</div>
        <span style="font-size:15px;font-weight:500;color:{TEXT};">ForecastIQ</span>
    </div>
    <hr style="margin:0 0 12px;"/>
    """, unsafe_allow_html=True)

    # ── Step 1: Data ──────────────────────────────────────────────
    st.markdown(f'<p style="font-size:10px;color:{TEXT3};text-transform:uppercase;letter-spacing:.06em;margin-bottom:6px;">① Data</p>', unsafe_allow_html=True)

    st.download_button(
        "↓ Get started (.zip)",
        data=_starter_kit_zip(),
        file_name="forecastiq_starter_kit.zip",
        mime="application/zip",
        use_container_width=True,
        help="Includes a step-by-step data guide (.docx) and a ready-to-use sample dataset (.csv).",
    )

    st.markdown('<div style="margin-top:8px;"></div>', unsafe_allow_html=True)
    uploaded_file = st.file_uploader("Upload CSV", type=["csv"], label_visibility="collapsed")

    use_sample = st.checkbox("Use built-in sample data", value=False)

    # Load data into session state
    if uploaded_file is not None:
        try:
            st.session_state.df_original = _read_csv_silent(uploaded_file)
            st.success("File loaded.")
        except Exception as exc:
            st.error(f"Could not read file: {exc}")
            st.session_state.df_original = None
    elif use_sample and st.session_state.df_original is None:
        st.session_state.df_original = generate_demo_data(num_skus=10, periods=104, granularity="Weekly")
        st.success("Sample data loaded.")

    df = st.session_state.df_original

    if df is None or df.empty:
        st.info("Upload a CSV or load the sample to begin.")
        st.stop()

    # Compute once after data load
    granularity    = detect_granularity(df, silent=True)
    quality        = data_quality_check(df, silent=True)
    effective_gran = "Weekly" if granularity == "Daily" else granularity
    effective_unit = horizon_unit_label(effective_gran)
    pareto_table, all_skus = pareto_analysis(df, top_n=quality["unique_skus"], silent=True)
    pareto_table = classify_abc(pareto_table)

    # ── Data quality cards ────────────────────────────────────────
    st.markdown('<hr style="margin:8px 0;"/>', unsafe_allow_html=True)
    q_cols = st.columns(2)
    q_cols[0].metric("Rows",    _fmt(quality["total_rows"]))
    q_cols[1].metric("SKUs",    _fmt(quality["unique_skus"]))
    q_cols2 = st.columns(2)
    q_cols2[0].metric("History", f"{quality['history_months']} mo")
    q_score = quality["quality_score"]
    q_cols2[1].metric("Quality",  f"{q_score}%",
                       delta="Good" if q_score >= 80 else ("Fair" if q_score >= 60 else "Poor"),
                       delta_color="normal" if q_score >= 80 else "inverse")

    if granularity == "Daily":
        st.caption("Daily data — forecasts auto-aggregate to weekly.")

    st.markdown('<hr style="margin:8px 0;"/>', unsafe_allow_html=True)

    # ── Step 2: Filters ───────────────────────────────────────────
    st.markdown(f"""
    <div style="display:flex;align-items:center;gap:6px;margin-bottom:6px;">
        <p style="font-size:10px;color:{TEXT3};text-transform:uppercase;letter-spacing:.06em;margin:0;">② Filters</p>
        <span style="font-size:9px;background:#14532d;color:#4ade80;border-radius:3px;padding:1px 5px;">RE-RUN TO APPLY</span>
    </div>
    """, unsafe_allow_html=True)

    industry = st.selectbox(
        "Industry",
        list(INDUSTRY_PROFILES.keys()),
        index=0,
        help=(
            "Sets the seasonality type and preferred model.\n\n"
            "ADDITIVE — General, Healthcare, Auto\n"
            "Seasonal swings stay the same size regardless of volume.\n"
            "Best when demand is steady with regular ups and downs.\n\n"
            "MULTIPLICATIVE — Retail, Semi\n"
            "Seasonal swings scale with volume level.\n"
            "Best when demand spikes are proportional to the base.\n\n"
            "Also acts as a tiebreaker when two models score equally."
        ),
        label_visibility="visible",
    )

    has_category = "Category" in df.columns
    has_location = "Location" in df.columns
    selected_category = None
    selected_location = None
    df_filtered = df.copy()

    if has_category:
        cats = ["All categories"] + sorted(df["Category"].dropna().unique().tolist())
        sel_cat = st.selectbox("Category", cats, label_visibility="visible")
        if sel_cat != "All categories":
            selected_category = sel_cat
            df_filtered = df_filtered[df_filtered["Category"] == sel_cat].copy()
    else:
        sel_cat = "All categories"

    if has_location:
        locs = ["All locations"] + sorted(df["Location"].dropna().unique().tolist())
        sel_loc = st.selectbox("Location", locs, label_visibility="visible")
        if sel_loc != "All locations":
            selected_location = sel_loc
            df_filtered = df_filtered[df_filtered["Location"] == sel_loc].copy()
    else:
        sel_loc = "All locations"

    # Recompute pareto after filter
    pareto_table, all_skus = pareto_analysis(df_filtered, top_n=df_filtered["SKU"].nunique(), silent=True)
    pareto_table = classify_abc(pareto_table)

    st.markdown('<hr style="margin:8px 0;"/>', unsafe_allow_html=True)

    # ── Step 3: SKU mode ──────────────────────────────────────────
    st.markdown(f'<p style="font-size:10px;color:{TEXT3};text-transform:uppercase;letter-spacing:.06em;margin-bottom:6px;">③ SKU mode</p>', unsafe_allow_html=True)

    sku_mode = st.radio(
        "SKU selection method",
        ["Pareto top N", "Manual pick"],
        index=0, horizontal=True,
        label_visibility="collapsed",
    )

    if sku_mode == "Pareto top N":
        pareto_mode = pareto_table["Pareto_Mode"].iloc[0] if len(pareto_table) else "volume"
        pareto_n = st.slider(
            f"Top N SKUs by {pareto_mode.lower()}",
            min_value=1, max_value=min(20, len(all_skus)),
            value=min(5, len(all_skus)), step=1,
        )
        selected_skus = all_skus[:pareto_n]
        st.caption(f"Top {pareto_n} SKUs by {pareto_mode.lower()}.")
    else:
        select_all = st.checkbox("Select all", value=False, key="select_all_skus")
        default_skus = all_skus if select_all else []
        selected_skus = st.multiselect(
            "SKUs to forecast",
            options=all_skus,
            default=default_skus,
            label_visibility="collapsed",
            placeholder="Select one or more SKUs…",
        )

    # Track whether user has made any active selection
    if selected_skus or sku_mode == "Pareto top N":
        st.session_state.user_has_selected = True

    st.markdown('<hr style="margin:8px 0;"/>', unsafe_allow_html=True)
    # ── Step 4: Horizon ───────────────────────────────────────────
    st.markdown(f'<p style="font-size:10px;color:{TEXT3};text-transform:uppercase;letter-spacing:.06em;margin-bottom:6px;">④ Horizon</p>', unsafe_allow_html=True)
    _max_h    = max_forecast_horizon(df_filtered, effective_gran)
    _default  = min(12, _max_h)
    horizon   = st.slider(
        f"Forecast horizon ({effective_unit})",
        min_value=4, max_value=_max_h, value=_default, step=1,
        label_visibility="visible",
    )
    st.caption(f"Max {_max_h} {effective_unit} based on history · cadence: {effective_gran.lower()}")

    if horizon > _max_h // 2:
        st.warning(f"Horizon over {_max_h // 2} {effective_unit} may reduce accuracy.")

    # ── Run button ────────────────────────────────────────────────
    st.markdown('<div style="margin-top:8px;"></div>', unsafe_allow_html=True)
    run_clicked = st.button("Run forecast", use_container_width=True)

    st.markdown(f'<p style="font-size:10px;color:{TEXT4};text-align:center;margin-top:8px;">Your data never leaves your browser. Anonymous usage stats logged.</p>', unsafe_allow_html=True)


# ── Mobile CSS ────────────────────────────────────────────────────────────────
st.markdown(f"""
<style>
@media (max-width: 768px) {{
    [data-testid="stSidebar"] {{
        width: 85vw !important;
        min-width: unset !important;
    }}
    .block-container {{
        padding-left: 0.75rem !important;
        padding-right: 0.75rem !important;
        padding-top: 1rem !important;
    }}
    [data-testid="stMetric"] {{
        padding: 8px 10px !important;
    }}
    [data-testid="stMetric"] [data-testid="stMetricValue"] {{
        font-size: 0.95rem !important;
    }}
}}
</style>
""", unsafe_allow_html=True)

# ── Main area ─────────────────────────────────────────────────────────────────
if df is None or df.empty:
    st.info("Load data to get started.")
    st.stop()

if not selected_skus:
    if not st.session_state.user_has_selected:
        st.markdown(f"""
        <div style="padding:48px 24px;text-align:center;color:{TEXT3};font-size:13px;line-height:1.8;">
            <div style="font-size:15px;color:{TEXT2};font-weight:500;margin-bottom:8px;">
                Ready to forecast
            </div>
            Use <strong style="color:{TEXT};">Pareto top N</strong> to auto-select your highest-value SKUs,
            or switch to <strong style="color:{TEXT};">Manual pick</strong> to choose specific ones.<br>
            Then click <strong style="color:{TEXT};">Run forecast</strong>.
        </div>
        """, unsafe_allow_html=True)
    else:
        st.info("Select at least one SKU in the sidebar, then click Run forecast.")
    st.stop()

# ── Run engine only when inputs change ───────────────────────────────────────
run_key = _run_key(selected_skus, horizon, industry, selected_location)

if run_clicked or (run_key != st.session_state.last_run_key and not st.session_state.forecast_results):
    with st.spinner("Running backtests and generating forecasts…"):
        results: Dict[str, dict] = {}
        for sku in selected_skus:
            try:
                results[sku] = forecast_sku(
                    df=df_filtered, sku=sku, granularity=granularity,
                    forecast_horizon=horizon, industry=industry,
                    n_splits=4, silent=True,
                    location=selected_location,
                )
            except Exception as exc:
                st.warning(f"Could not forecast SKU {sku}: {exc}")
        st.session_state.forecast_results = results
        st.session_state.last_run_key = run_key
        # Log usage — no file data, counts and settings only
        _log_run(
            skus_count  = len(selected_skus),
            horizon     = horizon,
            industry    = industry,
            granularity = granularity,
            sku_mode    = sku_mode,
        )

results = st.session_state.forecast_results

if not results:
    st.info("Click **Run forecast** in the sidebar to generate forecasts.")
    st.stop()

skipped    = {k: v for k, v in results.items() if v.get("skipped")}
forecasted = {k: v for k, v in results.items() if not v.get("skipped")}

if skipped:
    skipped_names = ", ".join(skipped.keys())
    st.caption(f"Skipped ({len(skipped)}): {skipped_names} — fewer than {MIN_DENSITY_RECORDS} data points.")

if not forecasted:
    st.warning("All selected SKUs were skipped due to insufficient data.")
    st.stop()

# ── KPI strip ────────────────────────────────────────────────────────────────
first_result  = next(iter(forecasted.values()))
first_fdf     = first_result["forecast_df"]
forecast_start = pd.to_datetime(first_fdf["Date"].iloc[0]).strftime("%b %Y")
forecast_end   = pd.to_datetime(first_fdf["Date"].iloc[-1]).strftime("%b %Y")

all_hist_vals = np.concatenate([
    v["historical_df"]["Quantity"].values for v in forecasted.values()
])
all_fc_vals = np.concatenate([
    v["forecast_df"]["Best_Model"].values for v in forecasted.values()
])
avg_monthly_demand = round(float(np.mean(all_fc_vals)), 1)

# Demand Trend — computed per SKU then averaged to avoid window size inflation
trend_pcts = []
for r in forecasted.values():
    hist_vals = r["historical_df"]["Quantity"].values.astype(float)
    fc_vals   = r["forecast_df"]["Best_Model"].values.astype(float)
    trend_pcts.append(compute_demand_trend(hist_vals, fc_vals))
trend_pct   = round(float(np.mean(trend_pcts)), 1) if trend_pcts else 0.0
trend_label = f"{trend_pct:+.1f}%"

kpi1, kpi2, kpi3 = st.columns(3)
kpi1.metric("Forecast period",     f"{forecast_start} – {forecast_end}")
kpi2.metric("Avg forecast demand", f"{avg_monthly_demand:.0f} units")
kpi3.metric("Demand trend",        trend_label)

st.markdown('<div style="margin-bottom:4px;"></div>', unsafe_allow_html=True)

# ── Chart area ────────────────────────────────────────────────────────────────
sku_list = list(forecasted.keys())

chart_col, _ = st.columns([1, 0.001])
with chart_col:
    with st.container():
        # Chart header row
        h_left, h_right = st.columns([2, 1])
        with h_left:
            focus_sku = st.selectbox(
                "SKU",
                options=sku_list,
                index=0,
                label_visibility="collapsed",
                key="sku_selector",
            )
        with h_right:
            focus = forecasted[focus_sku]
            model_options = [r["Method"] for r in focus["accuracy_results"]]
            best_method   = focus["best_method"]
            default_idx   = model_options.index(best_method) if best_method in model_options else 0

            active_model = st.selectbox(
                "Model",
                options=model_options,
                index=default_idx,
                format_func=lambda m: f"{m} ★" if m == best_method else m,
                label_visibility="collapsed",
                key="model_selector",
            )

        # Warn only if user picks a worse model
        st.caption("★ = recommended model")
        if active_model != best_method:
            best_mape  = focus["accuracy_results"][focus["best_index"]]["MAPE_%"]
            active_row = next((r for r in focus["accuracy_results"] if r["Method"] == active_model), None)
            if active_row and active_row["MAPE_%"] > best_mape:
                st.warning(f"{active_model} MAPE {active_row['MAPE_%']}% vs {best_method} {best_mape}% — recommended model has lower error.")

        # Chart
        st.plotly_chart(_build_chart(focus, active_model), use_container_width=True)

        # Export button — inline after chart
        export_df = _build_export(forecasted, pareto_table)
        if not export_df.empty:
            buf = io.StringIO()
            export_df.to_csv(buf, index=False)
            st.download_button(
                "↓ Export forecast CSV",
                data=buf.getvalue(),
                file_name="forecastiq_forecast.csv",
                mime="text/csv",
            )

st.markdown('<div style="margin-bottom:4px;"></div>', unsafe_allow_html=True)

# ── Forward planning table ────────────────────────────────────────────────────
fwd_df = _forward_table(forecasted, effective_gran, pareto_table)

if not fwd_df.empty:
    st.markdown(f"""
    <div style="background:{BG2};border-radius:10px;padding:1px 0 0;">
    """, unsafe_allow_html=True)

    # Sort by 3-period total descending, show top 10 by default
    try:
        fwd_sorted = fwd_df.copy()
        fwd_sorted["_sort"] = fwd_sorted["3-period total"].str.replace(",","").astype(float)
        fwd_sorted = fwd_sorted.sort_values("_sort", ascending=False).drop(columns=["_sort"])
    except Exception:
        fwd_sorted = fwd_df

    show_all = st.checkbox(f"Show all {len(fwd_sorted)} SKUs", value=False, key="show_all_fwd")
    display_df = fwd_sorted if show_all else fwd_sorted.head(10)

    st.dataframe(display_df, use_container_width=True, hide_index=True)

    # ── Model details (collapsed) ─────────────────────────────────
    with st.expander("▾ Model details — WAPE / RMSE / Folds"):
        focus_for_detail = forecasted[focus_sku]
        acc_df = pd.DataFrame(focus_for_detail["accuracy_results"])
        acc_df = acc_df[["Method", "MAPE_%", "RMSE", "Folds"]].rename(columns={"MAPE_%": "WAPE %"})
        best_m = focus_for_detail["best_method"]
        acc_df["Method"] = acc_df["Method"].apply(lambda m: f"{m} ★" if m == best_m else m)
        acc_df = acc_df.sort_values("WAPE %")
        st.dataframe(acc_df, use_container_width=True, hide_index=True)

    st.markdown("</div>", unsafe_allow_html=True)

# ── Footer ────────────────────────────────────────────────────────────────────
st.markdown(f"""
<div style="text-align:center;padding:1.5rem 0 0.5rem;color:{TEXT4};font-size:0.75rem;">
    Built by Preet Patel · ForecastIQ v2.0
</div>
""", unsafe_allow_html=True)
